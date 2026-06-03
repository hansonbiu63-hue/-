from __future__ import annotations

import csv
import json
import re
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\UnityProgram\test2")
OUT_DIR = ROOT / "deliverables" / "docx_draft"
REFERENCE_ROOT = Path(
    r"C:\Users\Lenovo\Desktop\交互组合仿真子系统源码＋APK\交互组合仿真子系统源码＋APK"
)
IMPLEMENTATION_PLAN = REFERENCE_ROOT / "04_Documents" / "交互组合仿真子系统技术项目实施方案.docx"
SOFTWARE_GUIDE = REFERENCE_ROOT / "04_Documents" / "交互组合仿真子系统.docx"
APK_DIR = REFERENCE_ROOT / "02_APK"
MODEL_DATA_DIR = REFERENCE_ROOT / "03_Model_Data"
TEST_DATA_DIR = Path(r"F:\TestData")
SYSTEM_NAME = "交互组合仿真子系统"
VERSION = "V1.0 草稿"
TODAY = datetime.now().strftime("%Y年%m月%d日")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"
WHITE = "FFFFFF"


def docx_text(path: Path) -> str:
    if not path.exists():
        return ""
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs: list[str] = []
        for para in root.findall(".//w:p", ns):
            text = "".join(t.text or "" for t in para.findall(".//w:t", ns)).strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)
    except Exception as exc:
        return f"[读取失败: {path}，{exc}]"


def safe_read(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="gbk", errors="ignore")
    except FileNotFoundError:
        return ""


def load_json(path: Path) -> dict:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def list_files(path: Path, pattern: str = "*") -> list[Path]:
    if not path.exists():
        return []
    return sorted(path.glob(pattern), key=lambda p: p.name.lower())


def first_csv_header(path: Path) -> list[str]:
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            return next(reader)
    except Exception:
        return []


def count_csv_rows(path: Path) -> int:
    try:
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            return max(sum(1 for _ in f) - 1, 0)
    except Exception:
        return 0


def collect_context() -> dict:
    stats = load_json(ROOT / "Assets" / "Resources" / "eye_transformer_stats.json")
    metrics = load_json(TEST_DATA_DIR / "OUT_DIR" / "metrics_test.json")
    manifest = load_json(ROOT / "Packages" / "manifest.json")
    readme = safe_read(ROOT / "README.md")
    project_version = safe_read(ROOT / "ProjectSettings" / "ProjectVersion.txt")
    plan_text = docx_text(IMPLEMENTATION_PLAN)
    guide_text = docx_text(SOFTWARE_GUIDE)

    script_files = list_files(ROOT / "Assets" / "Scripts", "**/*.cs")
    apks = list_files(ROOT, "*.apk") + list_files(APK_DIR, "*.apk")
    frame_csvs = sorted(TEST_DATA_DIR.glob("Frame_*.csv")) if TEST_DATA_DIR.exists() else []
    representative_csv = max(frame_csvs, key=lambda p: p.stat().st_size) if frame_csvs else None
    csv_header = first_csv_header(representative_csv) if representative_csv else []
    csv_rows = count_csv_rows(representative_csv) if representative_csv else 0

    package_deps = manifest.get("dependencies", {})
    key_packages = {
        "Unity Sentis": package_deps.get("com.unity.sentis", "未检出"),
        "XR Interaction Toolkit": package_deps.get("com.unity.xr.interaction.toolkit", "未检出"),
        "OpenXR": package_deps.get("com.unity.xr.openxr", "未检出"),
        "PICO XR": package_deps.get("com.unity.xr.picoxr", "未检出"),
        "TextMesh Pro": package_deps.get("com.unity.textmeshpro", "未检出"),
    }
    module_counter = Counter(p.parent.name for p in script_files)

    return {
        "plan_text": plan_text,
        "guide_text": guide_text,
        "readme": readme,
        "project_version": project_version.strip(),
        "stats": stats,
        "metrics": metrics,
        "key_packages": key_packages,
        "script_files": script_files,
        "module_counter": module_counter,
        "apks": apks,
        "frame_csvs": frame_csvs,
        "representative_csv": representative_csv,
        "csv_header": csv_header,
        "csv_rows": csv_rows,
        "model_data_files": list_files(MODEL_DATA_DIR, "*"),
        "videos": list_files(REFERENCE_ROOT, "*.mp4"),
        "logs": list_files(ROOT / "Logs", "*.log"),
    }


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, 9.5, bold, color or INK)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def configure_document(doc: Document, title: str, doc_type: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{SYSTEM_NAME} | {doc_type}")
    set_run_font(run, 9, False, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{VERSION} | {TODAY}")
    set_run_font(run, 9, False, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SYSTEM_NAME)
    set_run_font(run, 18, True, DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, 24, True, RGBColor(0, 0, 0))

    add_metadata_table(
        doc,
        [
            ("文档版本", VERSION),
            ("编制日期", TODAY),
            ("文档状态", "草稿，供内部审改"),
            ("适用对象", "项目管理、开发、测试、主试和验收人员"),
            ("资料来源", "项目实施方案、软件说明书、当前 Unity 工程、测试数据与交付包"),
        ],
        widths=[1800, 7560],
    )
    doc.add_page_break()
    add_revision_table(doc)
    doc.add_paragraph("目录").style = "Heading 1"
    add_para(doc, "此处为目录占位。正式定稿前可在 Word 中通过“引用 - 目录 - 更新目录”生成自动目录。")
    doc.add_page_break()


def add_revision_table(doc: Document):
    doc.add_paragraph("修订记录").style = "Heading 1"
    add_table(
        doc,
        ["版本", "日期", "修订说明", "修订人"],
        [[VERSION, TODAY, "初始草稿生成，依据现有方案书、工程和测试记录整理。", "Codex"]],
        widths=[1200, 1800, 5160, 1200],
    )


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            set_run_font(run)
        if not p.runs:
            set_run_font(p.add_run(item))
        else:
            p.runs[0].text = item


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        if not p.runs:
            set_run_font(p.add_run(item))
        else:
            p.runs[0].text = item
            for run in p.runs:
                set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None):
    if widths is None:
        widths = [int(9360 / len(headers)) for _ in headers]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=RGBColor(0, 0, 0))
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            shade_cell(cells[idx], WHITE if len(table.rows) % 2 else "FAFBFC")
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_metadata_table(doc: Document, rows: list[tuple[str, str]], widths: list[int]):
    add_table(doc, ["项目", "内容"], [[a, b] for a, b in rows], widths=widths)


def add_source_note(doc: Document, sources: list[str]):
    doc.add_paragraph("资料来源与证据记录").style = "Heading 2"
    add_bullets(doc, sources)


def add_common_appendix(doc: Document, ctx: dict, purpose: str):
    doc.add_paragraph("附录A 术语与缩略语").style = "Heading 1"
    add_table(doc, ["术语", "说明"], [
        ["AOI", "Area of Interest，本文指实验界面中 Control_1 至 Control_10 等可被注视命中的目标控件区域。"],
        ["Frame CSV", "系统按帧记录的眼动与任务状态数据文件，是数据检查、模型训练和测试报告的重要证据来源。"],
        ["Top-3 推荐", "算法按概率或置信度选出的前三个候选控件，用于界面显示和任务辅助判断。"],
        ["ONNX", "开放神经网络交换格式，当前用于将离线训练完成的 Transformer 模型部署至 Unity 端。"],
        ["stats JSON", "与 ONNX 模型配套的统计配置文件，记录特征顺序、均值方差、窗口长度、阈值和 AOI 标签规则。"],
        ["静态任务", "目标控件和任务条件相对固定，主要用于稳定注视模式采集和基础识别验证。"],
        ["动态任务", "目标或控件状态随时间变化，主要用于验证连续任务和复杂时序意图识别能力。"],
    ], widths=[1800, 7560])
    doc.add_paragraph("附录B 当前版本边界").style = "Heading 1"
    add_bullets(doc, [
        f"本文档为 {purpose} 草稿，版本为 {VERSION}，用于内部审改和交付材料整理。",
        "当前文档已按“交互组合仿真子系统”统一系统名称，但封面编号、密级、签署页、盖章页仍待单位模板确认。",
        "测试与调试内容引用现有工程和数据记录；无法从文件证明的最终验收结论均不作夸大表述。",
        "当前模型指标来自已有 metrics_test.json 和 Frame CSV 记录，后续正式验收应以完整被试数据和现场测试记录为准。",
        "文档中路径为编制阶段证据路径，归档时可按交付目录规则替换为相对路径或附件编号。",
    ])
    doc.add_paragraph("附录C 证据与文件清单").style = "Heading 1"
    add_table(doc, ["类别", "文件/目录", "用途"], [
        ["原始方案", str(IMPLEMENTATION_PLAN), "项目背景、需求、调研、技术路线、进度和收益评估依据。"],
        ["软件说明", str(SOFTWARE_GUIDE), "系统组成、设备操作、软件流程、数据采集和算法说明依据。"],
        ["当前工程", str(ROOT), "Unity 工程、脚本、包依赖、模型配置和构建记录依据。"],
        ["交付包", str(REFERENCE_ROOT), "源码、APK、模型数据、已有文档和演示视频依据。"],
        ["测试数据", str(TEST_DATA_DIR), "Frame CSV、模型指标和调试记录依据。"],
        ["代表性 CSV", str(ctx["representative_csv"] or "未检出"), "字段、行数和采集记录样例。"],
        ["模型指标", str(TEST_DATA_DIR / "OUT_DIR" / "metrics_test.json"), "阶段性模型评估指标。"],
    ], widths=[1600, 5200, 2560])


def compact_paths(paths: Iterable[Path], limit: int = 8) -> str:
    vals = [str(p) for p in paths]
    shown = vals[:limit]
    suffix = "" if len(vals) <= limit else f"；另有 {len(vals) - limit} 项"
    return "；".join(shown) + suffix if shown else "未检出"


def metric_text(ctx: dict) -> str:
    metrics = ctx["metrics"]
    if not metrics:
        return "未检出 metrics_test.json，测试指标待补充。"
    parts = []
    mapping = [
        ("subset_acc", "子集准确率"),
        ("hamming_loss", "汉明损失"),
        ("macro_f1", "Macro-F1"),
        ("precision_at_3", "Precision@3"),
        ("recall_at_3", "Recall@3"),
        ("best_threshold", "最佳阈值"),
    ]
    for key, label in mapping:
        if key in metrics:
            val = metrics[key]
            if isinstance(val, float):
                val = f"{val:.4f}"
            parts.append(f"{label}={val}")
    return "；".join(parts)


def common_overview(ctx: dict) -> list[str]:
    stats = ctx["stats"]
    return [
        f"系统定位为面向 VR/AR 环境的眼动意图识别实验平台，融合 PICO 4 Pro 眼动采集、Unity 世界空间 UI、AOI 命中检测、实时推荐反馈和离线模型训练。",
        f"当前工程版本依据 ProjectVersion 为 {ctx['project_version'].replace(chr(10), '；')}；主构建场景为 Assets/Scenes/DemoScene.unity。",
        f"模型侧采用离线训练、ONNX 导出、Unity Sentis 推理的部署链路。当前 stats 文件记录 window_frames={stats.get('window_frames', '未检出')}、feature_dim={stats.get('feature_dim', '未检出')}、aoi_count={stats.get('aoi_count', '未检出')}、top_k={stats.get('top_k', '未检出')}、threshold={stats.get('threshold', '未检出')}。",
        "系统算法路线保留规则法、惩罚加权递归贝叶斯法、随机森林与 Transformer 方案；当前实现重点为轻量化 Transformer 在线推荐链路。",
    ]


def project_fact_table(ctx: dict) -> list[list[str]]:
    stats = ctx["stats"]
    return [
        ["开发引擎", "Unity 2022.3.62f3c1", "来自 ProjectSettings/ProjectVersion.txt"],
        ["目标设备", "PICO 4 Pro / PICO XR", "来自方案书、说明书与 PICO 包依赖"],
        ["主场景", "Assets/Scenes/DemoScene.unity", "来自 EditorBuildSettings.asset"],
        ["XR 依赖", "OpenXR、XR Interaction Toolkit、PICO XR SDK", "来自 Packages/manifest.json"],
        ["推理依赖", "Unity Sentis + ONNX 模型", "来自 Packages/manifest.json 与模型文件"],
        ["模型输入", f"150 帧窗口，{stats.get('feature_dim', 47)} 维特征", "来自 eye_transformer_stats.json"],
        ["输出目标", f"{stats.get('aoi_count', 10)} 个 AOI 控件，多标签 Sigmoid / Top-3 推荐", "来自 stats 与推理脚本"],
        ["数据记录", "Frame CSV、轮次/任务字段、AOI 命中、瞳孔/睁眼/有效性字段", "来自 GazeDataLogger 与测试 CSV"],
    ]


def requirements_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}需求分析报告", "需求分析报告")
    doc.add_paragraph("第一章 项目背景").style = "Heading 1"
    for item in common_overview(ctx):
        add_para(doc, item)
    add_para(doc, "原始实施方案明确提出，系统面向增强现实与虚拟现实场景下的自然交互需求，通过融合眼动通道与显示通道，根据被试者眼动数据推测其当前意图，并支撑 AR/VR 信息推送。")
    doc.add_paragraph("第二章 建设目标").style = "Heading 1"
    add_bullets(doc, [
        "构建稳定可运行的 PICO 眼动交互实验平台，支持设备接入、任务发布、AOI 命中、数据记录和推荐反馈。",
        "形成静态任务与动态任务两类实验流程，并支撑标准化眼动-意图数据采集。",
        "集成规则法、概率法和神经网络方法，形成可解释、可平滑、可学习的多层意图识别能力。",
        "支持离线训练、ONNX 导出、Unity 端推理和结果评估闭环。",
    ])
    doc.add_paragraph("第三章 功能需求").style = "Heading 1"
    add_table(doc, ["需求项", "需求说明", "当前工程对应"], [
        ["意图推断", "根据眼动数据判断被试正在监控或意图寻找的控件信息。", "EyeTrackingManager、ControlItem、RecommendationFeedbackManager"],
        ["多算法并行", "至少包括规则、概率和神经网络三类方法。", "RealtimeIntentionRecommender、OnnxIntentionInference、RandomForestIntentionInference"],
        ["任务场景", "支持静态固定任务和动态连续任务。", "TaskManager、TaskData、TaskType"],
        ["数据采集", "稳定记录帧级眼动、AOI、任务状态和有效性字段。", "GazeDataLogger 与 Frame CSV"],
        ["结果展示", "以推荐、高亮或文本形式展示 Top 候选目标。", "ControlItem 推荐高亮、结果文本组件"],
        ["评估统计", "统计窗口预测结果、任务结果与模型指标。", "WindowPredictionEvaluator、metrics_test.json"],
    ], widths=[1800, 4200, 3360])
    doc.add_paragraph("第四章 性能需求").style = "Heading 1"
    add_bullets(doc, [
        "静态任务目标：规则法、概率法和神经网络方法的推断正确率目标均不低于 90%。",
        "动态任务目标：至少一种算法的推断正确率目标不低于 85%。",
        "数据集目标：总实验次数不少于 100 次，总时长不少于 20 小时，被试人数不少于 10 人。",
        "运行目标：以 50Hz 左右采样节奏采集帧级数据，并按 150 帧窗口组织在线推理输入。",
    ])
    doc.add_paragraph("第五章 集成与约束").style = "Heading 1"
    add_table(doc, ["类别", "约束/要求", "说明"], [
        ["设备约束", "PICO 4 Pro 与 Windows 主机协同运行。", "需保证头显、手柄、权限、USB/ADB 或安装流程正常。"],
        ["坐标约束", "眼动数据需映射至 Unity 世界空间 UI 平面和屏幕归一化坐标。", "字段包括 screen_x_norm、screen_y_norm、hit_ui_plane。"],
        ["模型约束", "ONNX 文件与 stats JSON 必须保持特征顺序一致。", "OnnxIntentionInference 对 stats 文件进行校验。"],
        ["数据约束", "原始数据、轮次数据和评估数据需可追溯。", "以被试、Session、轮次和时间戳命名。"],
    ], widths=[1500, 4200, 3660])
    add_source_note(doc, [
        f"实施方案：{IMPLEMENTATION_PLAN}",
        f"软件说明书：{SOFTWARE_GUIDE}",
        f"工程 README：{ROOT / 'README.md'}",
        f"模型配置：{ROOT / 'Assets/Resources/eye_transformer_stats.json'}",
    ])
    return doc


def design_doc(ctx: dict, title: str, report_style: bool = False):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}{title}", title)
    doc.add_paragraph("第一章 设计依据").style = "Heading 1"
    for item in common_overview(ctx):
        add_para(doc, item)
    if report_style:
        add_para(doc, "本报告偏向评审表达，重点说明方案选择、技术路线、实现状态、风险与收益；详细开发接口和操作细节分别由程序员手册、用户手册和测试类文档承接。")
    doc.add_paragraph("第二章 总体架构").style = "Heading 1"
    add_table(doc, ["层级", "职责", "主要组件"], [
        ["设备与接入层", "接入 PICO 4 Pro、手柄和 XR Loader，获取眼动与输入信号。", "PXR SDK、OpenXR、XR Interaction Toolkit"],
        ["数据采集层", "采集视线原点、方向、命中坐标、AOI、瞳孔、有效性等帧级数据。", "EyeTrackingManager、GazeDataLogger"],
        ["任务与交互层", "维护静态/动态任务、轮次、目标控件、确认操作与结果状态。", "TaskManager、TaskData、ControlItem、ConfirmInput"],
        ["算法推理层", "执行规则、贝叶斯、随机森林或 Transformer 推理，输出控件概率和推荐结果。", "RealtimeIntentionRecommender、OnnxIntentionInference、RandomForestIntentionInference"],
        ["UI 与反馈层", "显示任务提示、数据状态、推荐结果，并对候选控件进行高亮。", "RecommendationFeedbackManager、RecommendationTruthFilterDisplay"],
        ["评估与训练层", "离线训练、ONNX 导出、窗口评估、指标统计与模型部署。", "train_eye_transformer_modified.py、WindowPredictionEvaluator、metrics_test.json"],
    ], widths=[1700, 4300, 3360])
    doc.add_paragraph("第三章 工程事实表").style = "Heading 1"
    add_table(doc, ["项目", "当前值", "来源"], project_fact_table(ctx), widths=[1800, 3900, 3660])
    doc.add_paragraph("第四章 核心数据流").style = "Heading 1"
    add_numbered(doc, [
        "PICO 设备启动眼动追踪并完成必要权限检查。",
        "EyeTrackingManager 读取眼动状态、视线方向、瞳孔和睁眼状态。",
        "系统将视线投射到世界空间 UI 平面，计算 screen_x_norm、screen_y_norm 和 AOI 命中结果。",
        "GazeDataLogger 按帧写入 CSV，并同步记录任务模式、轮次、目标控件和阶段。",
        "规则法、贝叶斯法或 Transformer 读取滑动窗口数据，输出控件概率或 Top 候选。",
        "RecommendationFeedbackManager 对推荐结果进行统一选择和界面反馈。",
        "WindowPredictionEvaluator 将窗口预测与任务目标对齐，形成评估记录。",
    ])
    doc.add_paragraph("第五章 算法设计").style = "Heading 1"
    add_table(doc, ["算法", "设计定位", "输入/输出", "当前状态"], [
        ["规则法", "基于驻留时长、AOI 命中和阈值逻辑的快速可解释判断。", "输入帧级 AOI 与注视统计，输出候选控件。", "作为基础对照和在线推荐链路保留。"],
        ["惩罚加权递归贝叶斯", "通过先验、惩罚系数和逐帧后验更新平滑动态任务判断。", "输入注视、速度、瞳孔、AOI 特征，输出后验概率。", "作为概率类方法保留。"],
        ["随机森林", "基于滑动窗口统计特征的多输出分类器，训练快、解释性较好。", "输入窗口统计特征，输出各控件概率。", "作为补充对照方法保留。"],
        ["轻量化 Transformer", "以时序编码器建模 150 帧窗口内的眼动模式，适合复杂动态任务。", "输入 [1,150,47]，输出 10 维 Sigmoid 概率与 Top-3。", "当前神经网络主线，采用 ONNX/Sentis 部署。"],
    ], widths=[1800, 3300, 2600, 1660])
    doc.add_paragraph("第六章 风险与改进").style = "Heading 1"
    add_table(doc, ["风险", "影响", "应对措施"], [
        ["眼动追踪权限或设备状态异常", "无法获得有效帧，影响采集与推理。", "启动前检查 PICO 权限、设备连接、日志与校准状态。"],
        ["AOI 映射偏差", "推荐结果与用户实际关注区域不一致。", "通过预实验检查 screen_x_norm、hit_ui_plane 和 ControlItem 命中。"],
        ["数据样本不足或类别不均衡", "模型泛化能力不足。", "补齐正式被试数据，使用 FocalLoss、阈值寻优和分层验证。"],
        ["ONNX 与 stats 不匹配", "推理特征顺序错误。", "部署时将模型和 stats JSON 成对更新，并保留校验逻辑。"],
    ], widths=[2800, 3000, 3560])
    add_source_note(doc, [
        f"当前脚本目录：{ROOT / 'Assets/Scripts'}",
        f"模型配置：{ROOT / 'Assets/Resources/eye_transformer_stats.json'}",
        f"模型数据目录：{MODEL_DATA_DIR}",
        f"交付包 APK：{compact_paths(list_files(APK_DIR, '*.apk'))}",
    ])
    return doc


def programmer_manual(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}程序员手册", "程序员手册")
    doc.add_paragraph("第一章 开发环境").style = "Heading 1"
    add_table(doc, ["项", "说明"], [
        ["Unity 版本", ctx["project_version"].replace("\n", "；")],
        ["目标平台", "Android / PICO 4 Pro"],
        ["XR 包", "PICO XR、OpenXR、XR Interaction Toolkit"],
        ["推理包", f"Unity Sentis {ctx['key_packages'].get('Unity Sentis')}"],
        ["主场景", "Assets/Scenes/DemoScene.unity"],
        ["构建产物", compact_paths(ctx["apks"], limit=5)],
    ], widths=[2000, 7360])
    doc.add_paragraph("第二章 工程目录").style = "Heading 1"
    add_table(doc, ["目录/文件", "用途"], [
        ["Assets/Scripts/EyeTracking", "眼动采集、推理、推荐反馈、窗口评估。"],
        ["Assets/Scripts/Task", "静态/动态任务管理、轮次状态、任务目标。"],
        ["Assets/Scripts/Control", "控件 ID、AOI 控件状态、推荐高亮。"],
        ["Assets/Scripts/System", "运行时场景补齐与组件绑定。"],
        ["Assets/Resources/eye_transformer_stats.json", "Transformer 特征顺序、均值方差、阈值、窗口参数。"],
        ["Assets/Models/eye_transformer_finetuned.onnx", "Unity 端在线推理模型。"],
        ["Tools/train_eye_transformer_modified.py", "离线预训练、微调和 ONNX 导出脚本。"],
    ], widths=[3300, 6060])
    doc.add_paragraph("第三章 核心类说明").style = "Heading 1"
    add_table(doc, ["脚本", "职责", "维护注意事项"], [
        ["EyeTrackingManager", "调用 PICO API、处理追踪状态、坐标映射、AOI 命中与帧记录。", "改动坐标系逻辑后必须验证 screen_x_norm、hit_ui_plane 和 aoi_id。"],
        ["GazeDataLogger", "写入 Frame CSV，维护当前轮次、任务目标和阶段。", "新增字段需同步训练脚本和文档字段表。"],
        ["TaskManager", "维护实验模式、任务生命周期、成功/失败状态和事件回调。", "改动任务流程需同步测试说明和操作规程。"],
        ["OnnxIntentionInference", "加载 ONNX/stats，构造 150x47 输入窗口，输出多标签概率。", "模型和 stats 必须成对部署。"],
        ["RecommendationFeedbackManager", "统一多算法推荐结果并驱动控件高亮。", "新增算法源时需维护结果优先级与清理逻辑。"],
        ["WindowPredictionEvaluator", "将推理窗口与任务区间对齐，生成评估统计。", "指标口径需与测试报告保持一致。"],
    ], widths=[2500, 4300, 2560])
    doc.add_paragraph("第四章 数据格式").style = "Heading 1"
    add_para(doc, f"代表性 CSV：{ctx['representative_csv'] or '未检出'}；记录行数约 {ctx['csv_rows']}。")
    add_table(doc, ["字段组", "字段示例", "用途"], [
        ["任务标识", "subject_id, session_id, round_index, mode, phase, target_controls", "定位被试、Session、轮次、任务模式和标签。"],
        ["时间索引", "timestamp_ms, frame_idx", "用于滑动窗口切分、延迟评估和数据对齐。"],
        ["眼动几何", "gaze_origin_*, gaze_dir_*, gaze_point_*", "描述视线原点、方向和投射点。"],
        ["屏幕映射", "screen_x_norm, screen_y_norm, hit_ui_plane", "描述视线在 UI 平面的归一化位置和命中状态。"],
        ["AOI 标签", "aoi_id, aoi_valid, aoi_name", "用于控件命中和训练标签对齐。"],
        ["生理与有效性", "pupil_*, pupil_valid, left_openness, right_openness", "用于注意力、失锁和眨眼处理。"],
    ], widths=[1900, 3700, 3760])
    doc.add_paragraph("第五章 模型训练与部署").style = "Heading 1"
    add_numbered(doc, [
        "采集并整理 Unity Frame CSV，确认字段完整、任务目标有效。",
        "使用 GazeBaseVR 或已有数据进行预训练，形成初始权重。",
        "使用 Unity 采集数据进行微调，生成 ONNX 模型和 stats JSON。",
        "将 ONNX 部署至 Assets/Models，将 stats JSON 部署至 Assets/Resources。",
        "启动 DemoScene，检查 OnnxIntentionInference 日志，确认输入形状、特征顺序和 AOI 标签通过校验。",
        "运行窗口评估，整理 metrics_test.json 或评估 CSV。",
    ])
    add_source_note(doc, [
        f"脚本数量统计：{dict(ctx['module_counter'])}",
        f"训练脚本：{ROOT / 'Tools/train_eye_transformer_modified.py'}",
        f"交付包模型数据：{compact_paths(ctx['model_data_files'])}",
    ])
    return doc


def user_or_operator_doc(ctx: dict, title: str, mode: str):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}{title}", title)
    doc.add_paragraph("第一章 适用范围").style = "Heading 1"
    if mode == "user":
        add_para(doc, "本文档面向最终使用人员和实验执行人员，说明系统启动、设备准备、实验任务执行、推荐结果查看和数据导出流程。")
    elif mode == "sop":
        add_para(doc, "本文档面向主试与操作人员，按操作规程形式规定实验准备、启动、采集、结束和异常处理步骤，保证执行一致性。")
    else:
        add_para(doc, "本文档面向维护人员，说明系统安装部署、运行维护、模型与数据维护、日志检查和常见故障处理方法。")
    doc.add_paragraph("第二章 设备与环境准备").style = "Heading 1"
    add_table(doc, ["检查项", "要求", "确认方式"], [
        ["PICO 4 Pro", "头显电量充足、佩戴稳定、眼动权限可用。", "开机后检查设备状态和眼动追踪状态。"],
        ["手柄", "左右手柄已连接，按键响应正常。", "在系统界面中执行确认输入测试。"],
        ["Windows 主机", "可安装 APK 或运行 Unity 工程调试。", "检查 USB/ADB、PICO 连接和工程依赖。"],
        ["实验环境", "光线适中、空间安全、被试无明显疲劳不适。", "主试按实验准备清单确认。"],
        ["软件资源", "APK、模型文件、stats JSON、数据目录准备完整。", "核对交付包和工程目录。"],
    ], widths=[1800, 4200, 3360])
    doc.add_paragraph("第三章 标准流程").style = "Heading 1"
    if mode == "user":
        steps = [
            "启动 PICO 设备并确认头显、手柄处于正常连接状态。",
            "安装并启动 PicoVREyeTracking APK，或在 Unity 中打开 DemoScene 进行调试。",
            "按提示完成设备佩戴、中心校正和必要的眼动追踪检查。",
            "进入系统主界面，确认任务模式、轮次和目标控件信息。",
            "开始静态或动态任务，按任务提示观察指定控件。",
            "观察推荐结果、高亮控件和右侧数据面板。",
            "任务结束后确认 CSV 数据已保存，再退出系统。",
        ]
    elif mode == "sop":
        steps = [
            "实验前准备：检查设备、电量、空间、主机、APK、模型和数据目录。",
            "被试准备：说明实验目的和操作要求，确认身体状态，完成佩戴调节。",
            "预实验：不纳入正式记录，用于检查 AOI 命中、按钮反馈、计时和数据写入。",
            "正式采集：主试发布任务，被试执行观察和确认，系统记录帧级数据。",
            "轮次结束：检查 Frame CSV 是否生成，确认任务结果和异常备注。",
            "阶段休息：每轮后安排短休，避免疲劳影响眼动数据质量。",
            "实验结束：备份数据、记录日志、关闭应用和设备。",
        ]
    else:
        steps = [
            "部署检查：确认 APK、源码、模型文件、stats JSON 和 README 同步。",
            "运行检查：启动系统并检查 PICO 权限、XR Loader、Sentis 模型加载日志。",
            "数据维护：定期整理 Frame CSV、评估记录和异常轮次说明。",
            "模型维护：更新 ONNX 时同步更新 stats JSON，并保留旧版备份。",
            "日志维护：保留 Unity 日志、导入日志和现场调试记录。",
            "版本维护：记录 APK、模型、源码和文档版本对应关系。",
        ]
    add_numbered(doc, steps)
    doc.add_paragraph("第四章 任务说明").style = "Heading 1"
    add_table(doc, ["任务类型", "操作要点", "记录重点"], [
        ["静态任务", "被试在固定时间段内观察主试指定的目标控件。", "注视持续时间、AOI 命中、推荐目标、完成情况。"],
        ["动态任务", "控件状态或任务目标连续变化，被试在目标满足条件时确认。", "任务切换、确认时刻、窗口预测、动态推荐稳定性。"],
        ["预实验", "用于熟悉流程和检查设备，不作为正式评估结论。", "设备异常、AOI 偏差、按钮响应和数据写入状态。"],
    ], widths=[1800, 4200, 3360])
    doc.add_paragraph("第五章 异常处理").style = "Heading 1"
    add_table(doc, ["异常", "可能原因", "处理方式"], [
        ["无眼动数据", "权限未授权、设备未支持、SDK 启动失败。", "检查 PICO 权限、重启应用、查看 EyeTracking 日志。"],
        ["AOI 全部为空", "UI 平面未绑定、坐标映射偏差、控件列表缺失。", "检查 uiPlaneRect、controlItems、ControlItem ID。"],
        ["推荐结果异常", "模型文件或 stats 不匹配、阈值不合适、样本质量不足。", "检查 OnnxIntentionInference 校验日志和 stats 参数。"],
        ["CSV 未生成", "数据目录权限、任务未启动、Logger 未绑定。", "检查 GazeDataLogger、TaskManager 绑定和输出路径。"],
        ["被试不适", "佩戴时间过长或眩晕。", "立即暂停实验，记录异常并安排休息。"],
    ], widths=[2200, 3300, 3860])
    if mode == "maintenance":
        doc.add_paragraph("第六章 维护清单").style = "Heading 1"
        add_table(doc, ["维护对象", "频率", "内容"], [
            ["APK 与源码", "每次版本更新", "记录版本号、构建时间、变更说明和回退包。"],
            ["模型与 stats", "每次训练后", "成对归档，记录数据集、阈值、指标和导出参数。"],
            ["测试数据", "每次实验后", "备份原始 CSV，整理无效轮次和异常备注。"],
            ["设备", "每次实验前后", "检查头显、手柄、电池、镜片和连接状态。"],
        ], widths=[2200, 1800, 5360])
    add_source_note(doc, [
        f"APK 目录：{APK_DIR}",
        f"演示视频：{compact_paths(ctx['videos'])}",
        f"工程 README：{ROOT / 'README.md'}",
        f"Unity 日志目录：{ROOT / 'Logs'}",
    ])
    return doc


def test_outline_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}项目测试大纲", "项目测试大纲")
    doc.add_paragraph("第一章 测试目标").style = "Heading 1"
    add_bullets(doc, [
        "验证系统是否满足眼动采集、任务管理、AOI 命中、数据记录、推荐反馈和模型推理等核心功能。",
        "验证静态任务与动态任务流程是否可稳定执行。",
        "验证现有模型、CSV 数据和评估指标是否具备可追溯性。",
        "为测试说明、测试报告和验收材料提供统一范围与判定准则。",
    ])
    doc.add_paragraph("第二章 测试范围").style = "Heading 1"
    add_table(doc, ["测试类别", "测试内容", "输出记录"], [
        ["功能测试", "启动、任务发布、眼动采集、AOI 命中、推荐显示、数据导出。", "测试用例记录、截图/视频、CSV。"],
        ["接口测试", "PICO SDK、XR Loader、Sentis、ONNX/stats 加载。", "Unity 日志、调试记录。"],
        ["数据测试", "CSV 字段完整性、时间戳连续性、目标标签和 AOI 有效性。", "Frame CSV 检查表。"],
        ["算法测试", "规则、贝叶斯、随机森林/Transformer 输出与指标统计。", "metrics_test.json、窗口评估记录。"],
        ["部署测试", "APK 安装、设备连接、主场景运行。", "APK 版本与安装记录。"],
    ], widths=[1800, 4700, 2860])
    doc.add_paragraph("第三章 测试环境").style = "Heading 1"
    add_table(doc, ["环境项", "配置"], [
        ["Unity", ctx["project_version"].replace("\n", "；")],
        ["设备", "PICO 4 Pro、左右手柄、Windows 主机"],
        ["核心依赖", "PICO XR、OpenXR、XR Interaction Toolkit、Unity Sentis"],
        ["测试数据", f"{TEST_DATA_DIR}"],
        ["代表性 APK", compact_paths(ctx["apks"], limit=4)],
    ], widths=[2200, 7160])
    doc.add_paragraph("第四章 准入与准出").style = "Heading 1"
    add_table(doc, ["阶段", "条件"], [
        ["测试准入", "源码、APK、模型、stats JSON、测试数据目录和设备准备完整；主场景可启动。"],
        ["测试准出", "核心用例完成；缺陷有记录；测试报告引用的数据、日志和指标可追溯。"],
    ], widths=[1800, 7560])
    doc.add_paragraph("第五章 判定准则").style = "Heading 1"
    add_bullets(doc, [
        "功能用例按“通过/不通过/受限通过/未执行”记录。",
        "性能目标区分“验收目标”和“已有记录指标”，不将现有样本指标直接等同最终验收结果。",
        "数据类用例要求 CSV 字段齐全、关键字段非空、任务标签可对齐。",
        "模型类用例要求 ONNX 与 stats 成对加载，推理结果可以在界面或评估记录中观察。",
    ])
    add_source_note(doc, [
        f"测试数据目录：{TEST_DATA_DIR}",
        f"模型指标：{TEST_DATA_DIR / 'OUT_DIR/metrics_test.json'}",
        f"日志目录：{ROOT / 'Logs'}",
    ])
    return doc


def test_description_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}项目测试说明", "项目测试说明")
    doc.add_paragraph("第一章 用例设计说明").style = "Heading 1"
    add_para(doc, "本说明按功能、数据、算法、部署和异常处理分类给出测试用例。测试执行时应记录操作人员、设备、APK/源码版本、模型版本、数据文件和实际结果。")
    doc.add_paragraph("第二章 测试用例").style = "Heading 1"
    rows = [
        ["TC-01", "系统启动", "安装/启动 APK 或运行 DemoScene", "进入主界面，无关键错误日志", "功能"],
        ["TC-02", "眼动权限", "首次启动并触发 PICO 眼动权限检查", "权限已授权或出现可处理提示", "接口"],
        ["TC-03", "AOI 命中", "观察 10 个控件区域并查看状态文本/CSV", "aoi_id 与 aoi_name 出现有效值", "功能"],
        ["TC-04", "静态任务", "选择静态任务并完成一轮", "任务开始、目标控件、结束状态被记录", "功能"],
        ["TC-05", "动态任务", "选择动态任务并完成连续任务", "任务切换、确认时刻和结果状态被记录", "功能"],
        ["TC-06", "Frame CSV", "任务后检查输出 CSV", "字段包含 timestamp_ms、screen_x_norm、aoi_id、target_controls 等", "数据"],
        ["TC-07", "ONNX 加载", "启动含模型推理的场景", "模型和 stats 校验通过，输入维度为 150x47", "算法"],
        ["TC-08", "推荐显示", "在有效任务中观察推荐面板和控件高亮", "至少一种算法输出候选推荐", "功能"],
        ["TC-09", "窗口评估", "运行含任务标签的数据并检查评估结果", "产生窗口级预测记录或统计指标", "算法"],
        ["TC-10", "异常中断", "模拟设备断开或无效追踪", "系统记录警告并不中断文档化流程", "异常"],
    ]
    add_table(doc, ["编号", "名称", "步骤", "预期结果", "类别"], rows, widths=[1100, 1700, 3200, 2500, 860])
    doc.add_paragraph("第三章 测试数据说明").style = "Heading 1"
    add_para(doc, f"现有 Frame CSV 数量：{len(ctx['frame_csvs'])}；代表性文件：{ctx['representative_csv'] or '未检出'}；代表性文件记录行数约：{ctx['csv_rows']}。")
    if ctx["csv_header"]:
        add_para(doc, "代表性 CSV 字段：" + "、".join(ctx["csv_header"][:32]) + (" 等" if len(ctx["csv_header"]) > 32 else ""))
    doc.add_paragraph("第四章 记录表").style = "Heading 1"
    add_table(doc, ["用例编号", "执行日期", "执行人", "实际结果", "证据文件", "结论"], [
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
    ], widths=[1300, 1500, 1200, 2400, 2200, 760])
    add_source_note(doc, [
        f"测试 CSV：{compact_paths(ctx['frame_csvs'], limit=6)}",
        f"模型指标：{TEST_DATA_DIR / 'OUT_DIR/metrics_test.json'}",
        f"APK：{compact_paths(ctx['apks'], limit=6)}",
    ])
    return doc


def test_report_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}测试报告", "测试报告")
    doc.add_paragraph("第一章 测试概述").style = "Heading 1"
    add_para(doc, "本报告依据当前可用工程、APK、Frame CSV、模型配置和 metrics_test.json 整理。报告区分“已有记录结果”和“验收目标指标”，不将当前样本记录替代最终验收结论。")
    doc.add_paragraph("第二章 测试对象").style = "Heading 1"
    add_table(doc, ["对象", "记录"], [
        ["源码工程", str(ROOT)],
        ["交付包源码", str(REFERENCE_ROOT / "01_Source_UnityProgram")],
        ["APK", compact_paths(ctx["apks"], limit=8)],
        ["模型数据", compact_paths(ctx["model_data_files"], limit=6)],
        ["测试数据", compact_paths(ctx["frame_csvs"], limit=8)],
        ["Unity 日志", compact_paths(ctx["logs"], limit=6)],
    ], widths=[2200, 7160])
    doc.add_paragraph("第三章 数据记录统计").style = "Heading 1"
    add_table(doc, ["统计项", "结果"], [
        ["Frame CSV 文件数量", str(len(ctx["frame_csvs"]))],
        ["代表性 CSV", str(ctx["representative_csv"] or "未检出")],
        ["代表性 CSV 行数", str(ctx["csv_rows"])],
        ["代表性 CSV 字段数量", str(len(ctx["csv_header"]))],
        ["最新 APK/交付包 APK", compact_paths(ctx["apks"], limit=4)],
    ], widths=[2600, 6760])
    doc.add_paragraph("第四章 模型指标记录").style = "Heading 1"
    add_para(doc, metric_text(ctx))
    per_label = ctx["metrics"].get("per_label_f1") if ctx["metrics"] else None
    if per_label:
        add_table(doc, ["AOI", "F1"], [[str(i), f"{v:.4f}"] for i, v in enumerate(per_label)], widths=[1800, 7560])
    doc.add_paragraph("第五章 结论与遗留问题").style = "Heading 1"
    add_table(doc, ["项目", "结论"], [
        ["功能链路", "从现有工程和记录看，系统已具备 PICO 眼动采集、任务管理、AOI 命中、CSV 记录和模型推理链路。"],
        ["测试数据", "已存在多份 Frame CSV，可作为测试和模型微调记录来源。"],
        ["模型指标", "metrics_test.json 已记录当前测试集指标；当前指标用于阶段性分析，不作为最终验收性能结论。"],
        ["遗留问题", "需继续补齐正式被试规模、完整 Summary/Eval 记录、现场测试签字和最终验收统计。"],
    ], widths=[2200, 7160])
    add_source_note(doc, [
        f"指标文件：{TEST_DATA_DIR / 'OUT_DIR/metrics_test.json'}",
        f"代表性 CSV：{ctx['representative_csv'] or '未检出'}",
        f"交付包 APK：{APK_DIR}",
    ])
    return doc


def debug_report_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}调试报告", "调试报告")
    doc.add_paragraph("第一章 调试范围").style = "Heading 1"
    add_bullets(doc, [
        "PICO 眼动追踪支持检查、权限申请和启动流程。",
        "Unity 世界空间 UI 平面映射与 AOI 命中检测。",
        "静态/动态任务生命周期与手柄确认输入。",
        "Frame CSV 写入、任务标签更新和数据字段完整性。",
        "ONNX 模型加载、stats 校验、滑动窗口推理和推荐反馈。",
    ])
    doc.add_paragraph("第二章 关键调试项").style = "Heading 1"
    add_table(doc, ["调试项", "检查方法", "记录/结论"], [
        ["眼动权限", "查看 EyeTrackingManager 日志与 Android 权限状态。", "需确认 com.picovr.permission.EYE_TRACKING 已授权。"],
        ["追踪有效性", "观察 tracking_valid、hit_ui_plane、screen_x_norm。", "无效帧需保留标记，不应直接伪造有效数据。"],
        ["AOI ID", "观察 Control_1 至 Control_10 的命中状态。", "Control ID 需统一为 0-9 或经工具归一。"],
        ["CSV 写入", "任务结束后检查 Frame CSV。", f"现有代表性 CSV：{ctx['representative_csv'] or '未检出'}。"],
        ["模型加载", "启动 OnnxIntentionInference 并检查 stats。", "需匹配 150 帧、47 特征、10 AOI。"],
        ["推荐反馈", "观察推荐文本和控件高亮。", "推荐结果由 RecommendationFeedbackManager 统一发布。"],
    ], widths=[1900, 3900, 3560])
    doc.add_paragraph("第三章 日志与构建记录").style = "Heading 1"
    add_table(doc, ["类别", "记录"], [
        ["Unity 日志", compact_paths(ctx["logs"], limit=8)],
        ["APK 构建", compact_paths(ctx["apks"], limit=8)],
        ["模型文件", compact_paths(ctx["model_data_files"], limit=8)],
        ["演示视频", compact_paths(ctx["videos"], limit=4)],
    ], widths=[2200, 7160])
    doc.add_paragraph("第四章 常见问题处理").style = "Heading 1"
    add_table(doc, ["问题", "处理建议"], [
        ["StartEyeTracking 失败", "检查设备是否支持、权限是否授予、PICO SDK 配置是否正确。"],
        ["GetEyeTrackingData 失败", "检查追踪模式、设备佩戴和眼动服务状态。"],
        ["AOI 命中不稳定", "检查 UI 平面、控件 RectTransform、屏幕旋转和注视方向反转配置。"],
        ["Sentis 推理报错", "检查 ONNX opset、输入形状、stats 特征顺序和模型文件路径。"],
        ["CSV 字段缺失", "检查 GazeDataLogger 版本与训练脚本字段定义是否同步。"],
    ], widths=[2700, 6660])
    add_source_note(doc, [
        f"EyeTrackingManager：{ROOT / 'Assets/Scripts/EyeTracking/EyeTrackingManager.cs'}",
        f"OnnxIntentionInference：{ROOT / 'Assets/Scripts/EyeTracking/OnnxIntentionInference.cs'}",
        f"GazeDataLogger：{ROOT / 'Assets/Scripts/EyeTracking/GazeDataLogger.cs'}",
    ])
    return doc


def summary_doc(ctx: dict):
    doc = Document()
    configure_document(doc, f"{SYSTEM_NAME}项目研制总结", "项目研制总结")
    doc.add_paragraph("第一章 项目概述").style = "Heading 1"
    for item in common_overview(ctx):
        add_para(doc, item)
    doc.add_paragraph("第二章 研制过程").style = "Heading 1"
    add_table(doc, ["阶段", "主要工作", "成果"], [
        ["方案设计与评审", "完成实验任务、算法、数据集采集方案设计。", "形成技术项目实施方案。"],
        ["实验平台搭建", "完成 Unity/PICO 基础平台、眼动采集、任务界面和 AOI 命中。", "形成可运行原型和演示视频。"],
        ["中期实现与验证", "完成规则法、概率法、CSV 记录和初步数据采集。", "形成阶段数据和基础算法模块。"],
        ["系统集成与优化", "集成 Transformer ONNX 推理、推荐反馈和窗口评估。", "形成源码、APK、模型数据和文档材料。"],
    ], widths=[2100, 4300, 2960])
    doc.add_paragraph("第三章 主要成果").style = "Heading 1"
    add_bullets(doc, [
        "完成基于 PICO 4 Pro 的眼动采集和 Unity 世界空间 UI 映射。",
        "完成 10 个 AOI 控件的静态/动态任务框架和任务状态记录。",
        "完成 Frame CSV 数据记录链路，为后续训练和评估提供数据基础。",
        "完成规则法、贝叶斯法和神经网络方法的系统集成口径。",
        "完成轻量化 Transformer 的离线训练、ONNX 导出和 Unity Sentis 推理链路。",
        "形成源码、APK、模型数据、演示视频和本批 12 份文档草稿。",
    ])
    doc.add_paragraph("第四章 测试与指标").style = "Heading 1"
    add_para(doc, f"已有测试记录包括 {len(ctx['frame_csvs'])} 份 Frame CSV、metrics_test.json 和多版 APK。模型阶段性指标记录为：{metric_text(ctx)}。这些指标用于阶段总结和调优分析，最终验收仍需以正式测试报告和现场确认数据为准。")
    doc.add_paragraph("第五章 问题与建议").style = "Heading 1"
    add_table(doc, ["问题/风险", "建议"], [
        ["正式被试规模仍需与验收要求对齐", "按不少于 10 名被试、100 次实验和 20 小时总时长补齐记录。"],
        ["模型指标仍受样本规模和类别分布影响", "继续补充动态任务样本，进行阈值寻优和跨被试验证。"],
        ["文档仍为草稿", "待单位模板、编号、签署页和最终测试结论确认后定稿并导出 PDF。"],
        ["数据治理需要持续加强", "建立原始数据、预处理数据、训练数据和评估结果的版本映射。"],
    ], widths=[3500, 5860])
    add_source_note(doc, [
        f"实施方案：{IMPLEMENTATION_PLAN}",
        f"交付包根目录：{REFERENCE_ROOT}",
        f"当前工程：{ROOT}",
        f"测试数据：{TEST_DATA_DIR}",
    ])
    return doc


def save_doc(doc: Document, name: str):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / name
    doc.save(path)
    return path


def build_all():
    ctx = collect_context()
    docs = [
        ("交互组合仿真子系统需求分析报告.docx", requirements_doc(ctx)),
        ("交互组合仿真子系统软件设计方案.docx", design_doc(ctx, "软件设计方案")),
        ("交互组合仿真子系统设计方案报告.docx", design_doc(ctx, "设计方案报告", report_style=True)),
        ("交互组合仿真子系统程序员手册.docx", programmer_manual(ctx)),
        ("交互组合仿真子系统用户手册.docx", user_or_operator_doc(ctx, "用户手册", "user")),
        ("交互组合仿真子系统操作规程.docx", user_or_operator_doc(ctx, "操作规程", "sop")),
        ("交互组合仿真子系统使用维护说明书.docx", user_or_operator_doc(ctx, "使用维护说明书", "maintenance")),
        ("交互组合仿真子系统项目测试大纲.docx", test_outline_doc(ctx)),
        ("交互组合仿真子系统项目测试说明.docx", test_description_doc(ctx)),
        ("交互组合仿真子系统测试报告.docx", test_report_doc(ctx)),
        ("交互组合仿真子系统调试报告.docx", debug_report_doc(ctx)),
        ("交互组合仿真子系统项目研制总结.docx", summary_doc(ctx)),
    ]
    for name, doc in docs:
        purpose = name.removesuffix(".docx").replace(SYSTEM_NAME, "")
        add_common_appendix(doc, ctx, purpose)
    outputs = [save_doc(doc, name) for name, doc in docs]
    summary = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "output_dir": str(OUT_DIR),
        "documents": [str(p) for p in outputs],
        "source_documents": [str(IMPLEMENTATION_PLAN), str(SOFTWARE_GUIDE)],
        "frame_csv_count": len(ctx["frame_csvs"]),
        "representative_csv": str(ctx["representative_csv"] or ""),
        "metrics": ctx["metrics"],
        "apks": [str(p) for p in ctx["apks"]],
    }
    summary_path = OUT_DIR / "generation_summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return outputs


if __name__ == "__main__":
    for output in build_all():
        print(output)
